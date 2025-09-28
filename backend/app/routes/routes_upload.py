import shutil
import os
import logging
from fastapi import APIRouter, UploadFile, Depends, HTTPException, status, BackgroundTasks
from fastapi.concurrency import run_in_threadpool
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import update, delete # update import edildi
from app.db.session import get_db, async_session
from app.models.document import Document
from app.models.chunk import Chunk
from app.services.chunking import chunk_text
from app.services.embeddings import get_embedding_async 
from pypdf import PdfReader
from docx import Document as DocxDocument
import io

router = APIRouter()
UPLOAD_DIR = "uploaded_docs"

os.makedirs(UPLOAD_DIR, exist_ok=True)

logger = logging.getLogger(__name__)

# Dosya okuma fonksiyonu (Değişmedi, sadece context için dahil edildi)
def read_file_content(file: UploadFile) -> str:
    """
    Dosya uzantısına göre içeriği okur ve metin olarak döndürür.
    Desteklenen türler: .txt, .pdf, .docx
    """
    try:
        if file.filename.endswith(".pdf"):
            file.file.seek(0)
            pdf_reader = PdfReader(file.file)
            content = ""
            for page in pdf_reader.pages:
                content += page.extract_text() or ""
            return content
        elif file.filename.endswith(".docx"):
            file.file.seek(0)
            doc = DocxDocument(io.BytesIO(file.file.read()))
            content = ""
            for para in doc.paragraphs:
                content += para.text + "\n"
            return content
        elif file.filename.endswith(".txt"):
            file.file.seek(0)
            return file.file.read().decode("utf-8")
        else:
            raise HTTPException(
                status_code=status.HTTP_415_UNSUPPORTED_MEDIA_TYPE,
                detail=f"Desteklenmeyen dosya formatı: {file.filename.split('.')[-1]}. Sadece PDF, DOCX, TXT desteklenir."
            )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Dosya okuma hatası: {e}")
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Dosya içeriği okunamadı: {file.filename}. Hata: {str(e)}"
        )

# Arka planda çalışacak asenkron fonksiyon
async def process_document_and_save_chunks(document_id: int, content: str):
    """
    Chunk oluşturma, embedding ve veritabanına kaydetme işlemlerini yürütür.
    Durum güncellemelerini içerir.
    """
    # Arka plan görevi için ayrı bir oturum al
    async with async_session() as db:
        try:
            # 1. Durumu PROCESSING olarak güncelle
            await db.execute(
                update(Document)
                .where(Document.id == document_id)
                .values(status="PROCESSING")
            )
            await db.commit()
            logger.info(f"Doküman ID {document_id} durumu: PROCESSING")
            
            # 2. Chunk'lara ayır ve embedding oluştur
            chunks = chunk_text(content)
            if not chunks:
                raise ValueError("Dosya içeriği boş veya işlenemedi.")

            for idx, ch in enumerate(chunks, start=1):
                embedding = await get_embedding_async(ch) 
                new_chunk = Chunk(document_id=document_id, content=ch, embedding=embedding)
                db.add(new_chunk)
            
            # 3. Tüm Chunk'ları kaydet ve durumu READY olarak güncelle
            await db.execute(
                update(Document)
                .where(Document.id == document_id)
                .values(status="READY")
            )
            await db.commit()
            logger.info(f"Doküman ID {document_id} durumu: READY. {len(chunks)} parça başarıyla işlendi.")

        except Exception as e:
            await db.rollback()
            logger.error(f"Arka plan işleme hatası (doc_id={document_id}): {e}.")
            
            # 4. Hata durumunda durumu FAILED olarak güncelle ve tüm chunk'ları sil
            await db.execute(
                update(Document)
                .where(Document.id == document_id)
                .values(status="FAILED")
            )
            # Chunk'ların silinmesi için delete komutunu da kullanabiliriz (CASCADE yoksa)
            await db.execute(
                delete(Chunk).where(Chunk.document_id == document_id)
            )
            await db.commit()

# --- ANA ENDPOINT ---
@router.post("/upload", status_code=status.HTTP_201_CREATED)
async def upload_document(
    file: UploadFile, 
    background_tasks: BackgroundTasks,
    db: AsyncSession = Depends(get_db)
):
    """
    Yüklenen dosyayı kabul eder, PENDING durumuyla kaydını oluşturur ve 
    işlemi arka plana devreder.
    """
    
    # 1. Dosya içeriğini oku
    try:
        file_content = await run_in_threadpool(read_file_content, file) 
    except HTTPException as e:
        raise e
    except Exception as e:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail=str(e))
    
    
    # 2. Veritabanına yeni bir doküman kaydı oluştur (Durum varsayılan olarak PENDING)
    try:
        new_doc = Document(filename=file.filename)
        db.add(new_doc)
        await db.commit()
        await db.refresh(new_doc)
        logger.info(f"Doküman veritabanına eklendi: {new_doc.filename} (id={new_doc.id}, status=PENDING)")
    except Exception as e:
        await db.rollback()
        logger.error(f"Doküman veritabanına kaydedilemedi: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Dosya bilgisi veritabanına kaydedilirken bir hata oluştu."
        )
    
    # 3. İşleme görevini arka plana devret
    background_tasks.add_task(process_document_and_save_chunks, new_doc.id, file_content)

    return {
        "id": new_doc.id, 
        "filename": new_doc.filename, 
        "status": new_doc.status, # Frontend için durumu döndür
        "message": "Doküman başarıyla yüklendi. İşleme arka planda devam ediyor."
    }
